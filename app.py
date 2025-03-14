import os
import subprocess
import logging
import mimetypes
from flask import Flask, request, render_template, send_file, jsonify, redirect, url_for, make_response
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from flask_bcrypt import Bcrypt
from models import db, User
from audio_converter import download_ffmpeg, convert_audio
import traceback
import io
from PIL import Image

logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

app = Flask(__name__)
app.secret_key = 'replace_with_your_super_secret_key'

ALLOWED_AUDIO_EXTENSIONS = {'wav', 'mp3', 'ogg', 'flac', 'aac', 'm4a'}
ALLOWED_MIME_TYPES = {
    'audio/wav', 'audio/mpeg', 'audio/ogg', 'audio/flac', 
    'audio/aac', 'audio/x-m4a', 'audio/mp3'
}

ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
ALLOWED_IMAGE_MIME_TYPES = {
    'image/png', 'image/jpeg', 'image/gif', 'image/webp'
}

ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'docx', 'txt', 'md', 'html'}
ALLOWED_DOCUMENT_MIME_TYPES = {
    'application/pdf', 
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'text/plain', 
    'text/markdown', 
    'text/html'
}

# Initialize extensions
bcrypt = Bcrypt(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'

# Configure database
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///users.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db.init_app(app)

@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

@app.route('/', methods=['GET', 'POST'])
def index():
    return render_template('index.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user and bcrypt.check_password_hash(user.password, password):
            login_user(user)
            return redirect(url_for('index'))
        return render_template('login.html', error='Invalid username or password')
    return render_template('login.html')

@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')
        if User.query.filter_by(username=username).first():
            return render_template('signup.html', error='Username already exists')
        if User.query.filter_by(email=email).first():
            return render_template('signup.html', error='Email already registered')
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        user = User(username=username, email=email, password=hashed_password)
        db.session.add(user)
        db.session.commit()
        login_user(user)
        return redirect(url_for('index'))
    return render_template('signup.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index'))

@app.route('/convert-file', methods=['POST'])
def convert_file_route():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': 'No selected file'}), 400
        
    # Get original filename without extension
    original_filename = os.path.splitext(file.filename)[0]
    
    filename = file.filename.lower()
    input_format = filename.rsplit('.', 1)[-1] if '.' in filename else ''
    output_format = request.form.get('output_format', '').lower()
    
    # Get custom filename if provided, otherwise use original filename + "copy"
    custom_filename = request.form.get('custom_filename')
    if custom_filename and custom_filename.strip():
        output_filename = f"{custom_filename}.{output_format}"
    else:
        output_filename = f"{original_filename}_copy.{output_format}"
    
    # Determine file type and validate
    mime_type = file.content_type
    
    is_audio = (input_format in ALLOWED_AUDIO_EXTENSIONS and 
                output_format in ALLOWED_AUDIO_EXTENSIONS and 
                mime_type in ALLOWED_MIME_TYPES)
    
    is_image = (input_format in ALLOWED_IMAGE_EXTENSIONS and 
                output_format in ALLOWED_IMAGE_EXTENSIONS and 
                mime_type in ALLOWED_IMAGE_MIME_TYPES)
    
    is_document = (input_format in ALLOWED_DOCUMENT_EXTENSIONS and 
                  output_format in ALLOWED_DOCUMENT_EXTENSIONS and 
                  mime_type in ALLOWED_DOCUMENT_MIME_TYPES)
    
    if not (is_audio or is_image or is_document):
        return jsonify({
            'error': f'Invalid file type or conversion. Supported formats:\n' +
                    f'Audio: {", ".join(ALLOWED_AUDIO_EXTENSIONS)}\n' +
                    f'Images: {", ".join(ALLOWED_IMAGE_EXTENSIONS)}\n' +
                    f'Documents: {", ".join(ALLOWED_DOCUMENT_EXTENSIONS)}'
        }), 400
    
    # Create temp directory if it doesn't exist
    os.makedirs('temp', exist_ok=True)
    
    input_path = os.path.join('temp', f"temp_input.{input_format}")
    output_path = os.path.join('temp', f"temp_output.{output_format}")
    
    try:
        # Save uploaded file
        file.save(input_path)
        
        if is_audio:
            # Convert audio file
            logger.info(f"Converting audio from {input_format} to {output_format}")
            convert_audio(input_path, output_path)
        elif is_image:
            # Convert image file
            logger.info(f"Converting image from {input_format} to {output_format}")
            convert_image(input_path, output_path)
        else:
            # Convert document
            logger.info(f"Converting document from {input_format} to {output_format}")
            convert_document(input_path, output_path)
        
        # Read the converted file
        with open(output_path, 'rb') as f:
            file_data = f.read()
        
        # Clean up files
        try:
            os.remove(input_path)
            os.remove(output_path)
        except Exception as e:
            logger.error(f"Error cleaning up files: {str(e)}")
        
        # Send the response
        response = make_response(file_data)
        response.headers['Content-Type'] = mime_type
        response.headers['Content-Disposition'] = f'attachment; filename="{output_filename}"'
        return response
        
    except Exception as e:
        logger.error(f"Conversion error: {str(e)}")
        # Clean up files in case of error
        try:
            if os.path.exists(input_path):
                os.remove(input_path)
            if os.path.exists(output_path):
                os.remove(output_path)
        except Exception as cleanup_error:
            logger.error(f"Error cleaning up files: {str(cleanup_error)}")
        
        return jsonify({'error': 'An error occurred during conversion. Please try again.'}), 500

def convert_image(input_path: str, output_path: str):
    """Convert image using Pillow"""
    img = Image.open(input_path)
    
    # Convert RGBA to RGB if saving as JPEG
    if output_path.lower().endswith(('.jpg', '.jpeg')) and img.mode == 'RGBA':
        background = Image.new('RGB', img.size, (255, 255, 255))
        background.paste(img, mask=img.split()[3])
        img = background
    
    img.save(output_path, quality=95)

def convert_document(input_path: str, output_path: str):
    """Convert document using appropriate library based on input/output format"""
    input_format = input_path.split('.')[-1].lower()
    output_format = output_path.split('.')[-1].lower()
    
    try:
        if output_format == 'pdf':
            if input_format == 'docx':
                from docx2pdf import convert
                convert(input_path, output_path)
            elif input_format in ('txt', 'md', 'html'):
                import pdfkit
                with open(input_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if input_format == 'md':
                    import markdown
                    content = markdown.markdown(content)
                pdfkit.from_string(content, output_path)
                
        elif output_format == 'docx':
            from docx import Document
            doc = Document()
            
            if input_format == 'pdf':
                # Use pdf2docx for PDF to DOCX conversion
                from pdf2docx import Converter
                cv = Converter(input_path)
                cv.convert(output_path)
                cv.close()
            else:  # txt, md, html
                with open(input_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                if input_format == 'md':
                    import markdown
                    content = markdown.markdown(content)
                elif input_format == 'html':
                    from bs4 import BeautifulSoup
                    soup = BeautifulSoup(content, 'html.parser')
                    content = soup.get_text()
                doc.add_paragraph(content)
                doc.save(output_path)
            
        elif output_format in ('txt', 'md', 'html'):
            if input_format == 'pdf':
                from pdfminer.high_level import extract_text
                text = extract_text(input_path)
            elif input_format == 'docx':
                from docx import Document
                doc = Document(input_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            else:  # txt, md, html
                with open(input_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            
            # Convert to target format
            if output_format == 'html':
                if input_format == 'md':
                    import markdown
                    text = markdown.markdown(text)
                else:
                    text = f"<html><body><pre>{text}</pre></body></html>"
            elif output_format == 'md' and input_format == 'html':
                import html2text
                h = html2text.HTML2Text()
                text = h.handle(text)
                
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
                
    except Exception as e:
        logger.error(f"Document conversion error: {str(e)}")
        raise

if __name__ == '__main__':
    app.run(host='0.0.0.0', debug=True)
